import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import numpy as np
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import warnings
warnings.filterwarnings('ignore')

# matplotlib 백엔드 설정 (GUI 없이 이미지만 생성)
matplotlib.use('Agg')

# 한글 폰트 설정 (Windows 환경)
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False

def load_and_clean_data(file_path):
    """데이터 로드 및 전처리"""
    # CSV 파일 읽기
    df = pd.read_csv(file_path)
    
    print("=== 데이터 로드 중 ===")
    print(f"총 데이터 개수: {len(df)}개")
    
    # 데이터 정리
    # 1. 날짜 컬럼 변환
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
    
    # 2. 제품명과 카테고리 대소문자 정리
    df['ProductName'] = df['ProductName'].str.title()
    df['Category'] = df['Category'].str.title()
    df['Salesperson'] = df['Salesperson'].str.title()
    
    # 3. 빈 값이나 잘못된 데이터 제거
    df = df.dropna(subset=['Date', 'ProductID', 'Quantity', 'UnitPrice'])
    df = df[df['ProductID'] != 'P0000']
    df = df[df['Quantity'] > 0]
    df = df[df['UnitPrice'] > 0]
    
    # 4. TotalPrice 재계산
    df['TotalPrice'] = df['Quantity'] * df['UnitPrice']
    
    print(f"정리 후 데이터 개수: {len(df)}개")
    return df

def create_charts(df):
    """차트 생성 및 이미지 파일로 저장"""
    print("📊 차트 생성 중...")
    
    # 차트 파일들을 저장할 리스트
    chart_files = []
    
    try:
        # 1. 카테고리별 매출 파이차트
        category_sales = df.groupby('Category')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(8, 6))
        colors = plt.cm.Set3(np.linspace(0, 1, len(category_sales)))
        wedges, texts, autotexts = plt.pie(category_sales.values, labels=category_sales.index, 
                                          autopct='%1.1f%%', colors=colors, startangle=90)
        plt.title('카테고리별 매출 비율', fontsize=14, fontweight='bold', pad=20)
        
        # 범례 추가
        plt.legend(wedges, [f'{cat}\n{val:,.0f}원' for cat, val in category_sales.items()], 
                  title="카테고리별 매출액", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        
        plt.tight_layout()
        chart1_path = 'chart_category_pie.png'
        plt.savefig(chart1_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart1_path)
        
        # 2. 지역별 매출 막대차트
        region_sales = df.groupby('Region')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(10, 6))
        bars = plt.bar(region_sales.index, region_sales.values, color=['#FF9999', '#66B2FF', '#99FF99', '#FFCC99'])
        plt.title('지역별 매출액', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('지역', fontsize=12)
        plt.ylabel('매출액 (원)', fontsize=12)
        
        # 막대 위에 값 표시
        for bar, value in zip(bars, region_sales.values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(region_sales.values) * 0.01, 
                    f'{value:,.0f}원', ha='center', va='bottom', fontsize=10)
        
        plt.xticks(rotation=45)
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        chart2_path = 'chart_region_bar.png'
        plt.savefig(chart2_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart2_path)
        
        # 3. 일별 매출 추이 선 그래프
        daily_sales = df.groupby('Date')['TotalPrice'].sum().sort_index()
        
        plt.figure(figsize=(12, 6))
        plt.plot(daily_sales.index, daily_sales.values, marker='o', linewidth=2, markersize=4, color='#2E86AB')
        plt.title('일별 매출 추이', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('날짜', fontsize=12)
        plt.ylabel('매출액 (원)', fontsize=12)
        
        # 최고/최저점 표시
        max_idx = daily_sales.idxmax()
        min_idx = daily_sales.idxmin()
        plt.annotate(f'최고: {daily_sales[max_idx]:,.0f}원', 
                    xy=(max_idx, daily_sales[max_idx]), xytext=(10, 10),
                    textcoords='offset points', ha='left',
                    bbox=dict(boxstyle='round,pad=0.3', fc='yellow', alpha=0.7),
                    arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0'))
        
        plt.xticks(rotation=45)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        chart3_path = 'chart_daily_trend.png'
        plt.savefig(chart3_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart3_path)
        
        # 4. 베스트셀러 제품 TOP 5 막대차트
        top5_products = df.groupby('ProductName')['TotalPrice'].sum().sort_values(ascending=False).head(5)
        
        plt.figure(figsize=(10, 6))
        bars = plt.barh(range(len(top5_products)), top5_products.values, color='#FFB347')
        plt.title('베스트셀러 제품 TOP 5 (매출액 기준)', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('매출액 (원)', fontsize=12)
        plt.ylabel('제품명', fontsize=12)
        
        # y축 레이블 설정
        plt.yticks(range(len(top5_products)), top5_products.index)
        
        # 막대 끝에 값 표시
        for i, (bar, value) in enumerate(zip(bars, top5_products.values)):
            plt.text(bar.get_width() + max(top5_products.values) * 0.01, bar.get_y() + bar.get_height()/2, 
                    f'{value:,.0f}원', ha='left', va='center', fontsize=10)
        
        plt.grid(axis='x', alpha=0.3)
        plt.tight_layout()
        chart4_path = 'chart_top_products.png'
        plt.savefig(chart4_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart4_path)
        
        # 5. 영업사원별 성과 비교 차트
        df_clean = df[df['Salesperson'].notna() & (df['Salesperson'] != '')]
        salesperson_sales = df_clean.groupby('Salesperson')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(10, 6))
        bars = plt.bar(salesperson_sales.index, salesperson_sales.values, color='#98FB98')
        plt.title('영업사원별 매출 성과', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('영업사원', fontsize=12)
        plt.ylabel('매출액 (원)', fontsize=12)
        
        # 막대 위에 값 표시
        for bar, value in zip(bars, salesperson_sales.values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(salesperson_sales.values) * 0.01, 
                    f'{value:,.0f}원', ha='center', va='bottom', fontsize=9, rotation=0)
        
        plt.xticks(rotation=45)
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        chart5_path = 'chart_salesperson.png'
        plt.savefig(chart5_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart5_path)
        
        print(f"✅ {len(chart_files)}개의 차트가 생성되었습니다.")
        return chart_files
        
    except Exception as e:
        print(f"❌ 차트 생성 중 오류 발생: {e}")
        return []

def generate_word_report(df):
    """워드 파일(.docx) 보고서 생성"""
    print("="*30)
    print("📄 Word 보고서 생성 중...")
    print("="*30)
    
    try:
        # 차트 생성
        chart_files = create_charts(df)
        
        # 새 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading('판매 데이터 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성 일자 추가
        doc.add_paragraph(f"보고서 생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        doc.add_paragraph("")  # 빈 줄
        
        # === 1. 요약 통계 ===
        doc.add_heading('📊 요약 통계', level=1)
        
        # 요약 통계 표 생성
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        # 표 헤더
        summary_table.cell(0, 0).text = '구분'
        summary_table.cell(0, 1).text = '값'
        
        # 요약 데이터
        total_sales = df['TotalPrice'].sum()
        total_quantity = df['Quantity'].sum()
        avg_order_value = df['TotalPrice'].mean()
        unique_products = df['ProductID'].nunique()
        date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
        
        summary_data = [
            ['분석 기간', date_range],
            ['총 매출액', f"{total_sales:,.0f}원"],
            ['총 판매 수량', f"{total_quantity:,}개"],
            ['평균 주문 금액', f"{avg_order_value:,.0f}원"],
            ['판매된 제품 종류', f"{unique_products}개"]
        ]
        
        for i, (category, value) in enumerate(summary_data, 1):
            summary_table.cell(i, 0).text = category
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 2. 카테고리별 분석 ===
        doc.add_heading('📊 카테고리별 판매 분석', level=1)
        
        # 카테고리별 데이터 생성
        category_sales = df.groupby('Category').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum',
            'ProductID': 'nunique'
        }).round(2)
        category_sales.columns = ['총 매출액', '총 수량', '제품 종류 수']
        category_sales = category_sales.sort_values('총 매출액', ascending=False)
        
        # 카테고리별 분석 표
        cat_table = doc.add_table(rows=len(category_sales)+1, cols=4)
        cat_table.style = 'Table Grid'
        
        # 표 헤더
        cat_headers = ['카테고리', '총 매출액 (원)', '총 수량 (개)', '제품 종류 수']
        for i, header in enumerate(cat_headers):
            cat_table.cell(0, i).text = header
        
        # 카테고리 데이터
        for i, (category, row) in enumerate(category_sales.iterrows(), 1):
            cat_table.cell(i, 0).text = category
            cat_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            cat_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
            cat_table.cell(i, 3).text = str(row['제품 종류 수'])
        
        # 카테고리별 매출 파이차트 삽입
        if len(chart_files) > 0 and os.path.exists(chart_files[0]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 카테고리별 매출 비율 차트", style='Heading 2')
            doc.add_picture(chart_files[0], width=Inches(6))
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 3. 지역별 분석 ===
        doc.add_heading('🌍 지역별 판매 분석', level=1)
        
        # 지역별 데이터 생성
        region_sales = df.groupby('Region').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).round(2)
        region_sales.columns = ['총 매출액', '총 수량']
        region_sales = region_sales.sort_values('총 매출액', ascending=False)
        
        # 지역별 분석 표
        region_table = doc.add_table(rows=len(region_sales)+1, cols=3)
        region_table.style = 'Table Grid'
        
        # 표 헤더
        region_headers = ['지역', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(region_headers):
            region_table.cell(0, i).text = header
        
        # 지역 데이터
        for i, (region, row) in enumerate(region_sales.iterrows(), 1):
            region_table.cell(i, 0).text = region
            region_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            region_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
        
        # 지역별 매출 막대차트 삽입
        if len(chart_files) > 1 and os.path.exists(chart_files[1]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 지역별 매출 비교 차트", style='Heading 2')
            doc.add_picture(chart_files[1], width=Inches(6))
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 4. 베스트셀러 제품 ===
        doc.add_heading('🏆 베스트셀러 제품 TOP 10', level=1)
        
        # 베스트셀러 제품 데이터
        top_products = df.groupby(['ProductID', 'ProductName']).agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).sort_values('TotalPrice', ascending=False).head(10)
        
        # 베스트셀러 표
        product_table = doc.add_table(rows=len(top_products)+1, cols=4)
        product_table.style = 'Table Grid'
        
        # 표 헤더
        product_headers = ['순위', '제품명', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(product_headers):
            product_table.cell(0, i).text = header
        
        # 제품 데이터
        for i, ((product_id, product_name), row) in enumerate(top_products.iterrows(), 1):
            product_table.cell(i, 0).text = str(i)
            product_table.cell(i, 1).text = product_name
            product_table.cell(i, 2).text = f"{row['TotalPrice']:,.0f}"
            product_table.cell(i, 3).text = f"{row['Quantity']:,.0f}"
        
        # 베스트셀러 제품 차트 삽입
        if len(chart_files) > 3 and os.path.exists(chart_files[3]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 베스트셀러 제품 TOP 5 차트", style='Heading 2')
            doc.add_picture(chart_files[3], width=Inches(6))
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 5. 영업사원별 분석 ===
        doc.add_heading('👤 영업사원별 판매 성과', level=1)
        
        # 영업사원별 데이터 생성
        df_clean = df[df['Salesperson'].notna() & (df['Salesperson'] != '')]
        salesperson_data = df_clean.groupby('Salesperson').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum',
            'Date': 'count'
        }).sort_values('TotalPrice', ascending=False)
        
        # 영업사원 분석 표
        sales_table = doc.add_table(rows=len(salesperson_data)+1, cols=4)
        sales_table.style = 'Table Grid'
        
        # 표 헤더
        sales_headers = ['영업사원', '총 매출액 (원)', '총 수량 (개)', '거래 횟수']
        for i, header in enumerate(sales_headers):
            sales_table.cell(0, i).text = header
        
        # 영업사원 데이터
        for i, (salesperson, row) in enumerate(salesperson_data.iterrows(), 1):
            sales_table.cell(i, 0).text = salesperson
            sales_table.cell(i, 1).text = f"{row['TotalPrice']:,.0f}"
            sales_table.cell(i, 2).text = f"{row['Quantity']:,.0f}"
            sales_table.cell(i, 3).text = str(row['Date'])
        
        # 영업사원별 성과 차트 삽입
        if len(chart_files) > 4 and os.path.exists(chart_files[4]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 영업사원별 매출 성과 차트", style='Heading 2')
            doc.add_picture(chart_files[4], width=Inches(6))
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 6. 일별 매출 추이 분석 ===
        doc.add_heading('📈 일별 매출 추이 분석', level=1)
        
        # 일별 데이터 생성
        daily_sales = df.groupby('Date').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).round(2)
        daily_sales.columns = ['일별 매출액', '일별 수량']
        
        # 일별 통계
        max_sales_day = daily_sales['일별 매출액'].idxmax()
        min_sales_day = daily_sales['일별 매출액'].idxmin()
        avg_daily_sales = daily_sales['일별 매출액'].mean()
        
        daily_stats = doc.add_paragraph()
        daily_stats.add_run(f"• 최고 매출일: {max_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[max_sales_day, '일별 매출액']:,.0f}원\n")
        daily_stats.add_run(f"• 최저 매출일: {min_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[min_sales_day, '일별 매출액']:,.0f}원\n")
        daily_stats.add_run(f"• 일평균 매출액: {avg_daily_sales:,.0f}원")
        
        # 일별 매출 추이 차트 삽입
        if len(chart_files) > 2 and os.path.exists(chart_files[2]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 일별 매출 추이 차트", style='Heading 2')
            doc.add_picture(chart_files[2], width=Inches(7))
        
        # 일별 매출 TOP 10
        doc.add_paragraph("")
        doc.add_heading('일별 매출 TOP 10', level=2)
        
        top10_days = daily_sales.sort_values('일별 매출액', ascending=False).head(10)
        
        daily_table = doc.add_table(rows=len(top10_days)+1, cols=3)
        daily_table.style = 'Table Grid'
        
        # 표 헤더
        daily_headers = ['순위', '날짜', '매출액 (원)']
        for i, header in enumerate(daily_headers):
            daily_table.cell(0, i).text = header
        
        # 일별 데이터
        for i, (date, row) in enumerate(top10_days.iterrows(), 1):
            daily_table.cell(i, 0).text = str(i)
            daily_table.cell(i, 1).text = date.strftime('%Y-%m-%d')
            daily_table.cell(i, 2).text = f"{row['일별 매출액']:,.0f}"
        
        # === 7. 결론 및 제안사항 ===
        doc.add_heading('💡 결론 및 제안사항', level=1)
        
        # 자동 생성된 인사이트
        top_category = category_sales.index[0]
        top_region = region_sales.index[0]
        top_salesperson = salesperson_data.index[0]
        
        conclusions = doc.add_paragraph()
        conclusions.add_run(f"1. 주요 성과 카테고리: '{top_category}' 카테고리가 전체 매출의 주요 부분을 차지하고 있습니다.\n\n")
        conclusions.add_run(f"2. 핵심 지역: '{top_region}' 지역이 가장 높은 매출을 기록했습니다.\n\n")
        conclusions.add_run(f"3. 우수 영업사원: '{top_salesperson}' 사원이 최고 성과를 달성했습니다.\n\n")
        conclusions.add_run(f"4. 일평균 매출: {avg_daily_sales:,.0f}원으로, 지속적인 매출 관리가 필요합니다.\n\n")
        conclusions.add_run("5. 제안사항: 상위 성과 카테고리와 지역에 대한 마케팅 투자 확대를 고려해보시기 바랍니다.")
        
        # 문서 저장
        doc.save('sales_analysis_report.docx')
        print("✅ Word 보고서가 'sales_analysis_report.docx' 파일로 저장되었습니다.")
        
        # 임시 차트 파일들 정리 (선택적)
        # for chart_file in chart_files:
        #     if os.path.exists(chart_file):
        #         os.remove(chart_file)
        
        return True
        
    except Exception as e:
        print(f"❌ Word 파일 생성 중 오류 발생: {e}")
        return False

def send_email_with_report(docx_file_path, recipient_emails, sender_email=None, sender_password=None, 
                          smtp_server="smtp.gmail.com", smtp_port=587):
    """
    워드 보고서를 첨부하여 이메일 전송
    
    Args:
        docx_file_path (str): 전송할 워드 파일 경로
        recipient_emails (list): 받는 사람 이메일 주소 리스트
        sender_email (str): 보내는 사람 이메일 주소
        sender_password (str): 보내는 사람 이메일 비밀번호 (앱 비밀번호 권장)
        smtp_server (str): SMTP 서버 주소 (기본값: Gmail)
        smtp_port (int): SMTP 포트 번호 (기본값: 587)
    
    Returns:
        bool: 전송 성공 여부
    """
    print("="*30)
    print("📧 이메일 전송 중...")
    print("="*30)
    
    try:
        # 이메일 정보가 없는 경우 사용자 입력 받기
        if not sender_email:
            sender_email = input("보내는 사람 이메일 주소를 입력하세요: ")
        
        if not sender_password:
            import getpass
            sender_password = getpass.getpass("이메일 비밀번호(앱 비밀번호)를 입력하세요: ")
        
        # 이메일 메시지 생성
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = ", ".join(recipient_emails)
        msg['Subject'] = f"판매 데이터 분석 보고서 - {datetime.now().strftime('%Y년 %m월 %d일')}"
        
        # 이메일 본문 작성
        body = f"""
안녕하세요,

첨부된 파일은 {datetime.now().strftime('%Y년 %m월 %d일')}에 생성된 판매 데이터 분석 보고서입니다.

📊 보고서 내용:
• 요약 통계
• 카테고리별 판매 분석
• 지역별 판매 분석  
• 베스트셀러 제품 TOP 10
• 영업사원별 판매 성과
• 일별 매출 추이 분석
• 시각화 차트

보고서를 검토해 주시기 바랍니다.

감사합니다.

---
자동 생성된 보고서입니다.
"""
        
        # 본문을 이메일에 추가
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        # 워드 파일 첨부
        if os.path.exists(docx_file_path):
            with open(docx_file_path, "rb") as attachment:
                # MIMEBase 객체 생성
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
            
            # 파일을 base64로 인코딩
            encoders.encode_base64(part)
            
            # 헤더 추가
            filename = os.path.basename(docx_file_path)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {filename}',
            )
            
            # 이메일에 파일 첨부
            msg.attach(part)
            print(f"✅ 파일 '{filename}' 첨부 완료")
        else:
            print(f"❌ 파일을 찾을 수 없습니다: {docx_file_path}")
            return False
        
        # SMTP 서버 연결 및 이메일 전송
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()  # TLS 보안 연결 시작
        server.login(sender_email, sender_password)
        
        # 이메일 전송
        text = msg.as_string()
        server.sendmail(sender_email, recipient_emails, text)
        server.quit()
        
        print(f"✅ 이메일이 성공적으로 전송되었습니다!")
        print(f"📧 받는 사람: {', '.join(recipient_emails)}")
        
        return True
        
    except smtplib.SMTPAuthenticationError:
        print("❌ 이메일 인증에 실패했습니다.")
        print("💡 Gmail의 경우 '앱 비밀번호'를 사용해야 합니다.")
        print("   1. Google 계정 설정 > 보안 > 2단계 인증 활성화")
        print("   2. 앱 비밀번호 생성 후 해당 비밀번호 사용")
        return False
        
    except smtplib.SMTPException as e:
        print(f"❌ SMTP 오류가 발생했습니다: {e}")
        return False
        
    except Exception as e:
        print(f"❌ 이메일 전송 중 오류가 발생했습니다: {e}")
        return False

def send_email_interactive():
    """
    대화형 이메일 전송 함수
    사용자로부터 이메일 정보를 입력받아 보고서를 전송
    """
    print("="*40)
    print("📧 이메일 전송 설정")
    print("="*40)
    
    # 워드 파일 경로 확인
    docx_file = 'sales_analysis_report.docx'
    if not os.path.exists(docx_file):
        print(f"❌ 보고서 파일을 찾을 수 없습니다: {docx_file}")
        print("먼저 보고서를 생성해주세요.")
        return False
    
    try:
        # 받는 사람 이메일 주소 입력
        print("\n📮 받는 사람 이메일 주소를 입력하세요.")
        print("(여러 명에게 보낼 경우 쉼표로 구분: email1@gmail.com, email2@gmail.com)")
        recipients_input = input("받는 사람: ")
        recipient_emails = [email.strip() for email in recipients_input.split(',')]
        
        # 이메일 주소 형식 간단 검증
        for email in recipient_emails:
            if '@' not in email or '.' not in email:
                print(f"❌ 잘못된 이메일 형식: {email}")
                return False
        
        print(f"📧 {len(recipient_emails)}명에게 전송 예정")
        
        # SMTP 서버 설정 선택
        print("\n📡 이메일 서비스를 선택하세요:")
        print("1. Gmail (기본값)")
        print("2. Outlook/Hotmail") 
        print("3. 직접 입력")
        
        choice = input("선택 (1-3, 엔터시 Gmail): ").strip()
        
        if choice == "2":
            smtp_server = "smtp-mail.outlook.com"
            smtp_port = 587
        elif choice == "3":
            smtp_server = input("SMTP 서버 주소: ")
            smtp_port = int(input("SMTP 포트 번호: "))
        else:  # 기본값 또는 1
            smtp_server = "smtp.gmail.com"
            smtp_port = 587
        
        print(f"📡 SMTP 서버: {smtp_server}:{smtp_port}")
        
        # 이메일 전송 실행
        success = send_email_with_report(
            docx_file_path=docx_file,
            recipient_emails=recipient_emails,
            smtp_server=smtp_server,
            smtp_port=smtp_port
        )
        
        if success:
            print("\n🎉 이메일 전송이 완료되었습니다!")
        else:
            print("\n❌ 이메일 전송에 실패했습니다.")
            
        return success
        
    except KeyboardInterrupt:
        print("\n❌ 사용자가 취소했습니다.")
        return False
    except Exception as e:
        print(f"\n❌ 오류가 발생했습니다: {e}")
        return False

def main():
    """메인 함수"""
    print("🚀 판매 데이터 Word 보고서 생성을 시작합니다...\n")
    
    try:
        # 데이터 로드 및 전처리
        df = load_and_clean_data('cicd_data.csv')
        
        # Word 보고서 생성
        if generate_word_report(df):
            print("\n" + "="*50)
            print("✅ Word 보고서 생성이 완료되었습니다!")
            print("📄 생성된 파일: sales_analysis_report.docx")
            print("="*50)
            
            # 이메일 전송 여부 확인
            print("\n📧 생성된 보고서를 이메일로 전송하시겠습니까?")
            send_choice = input("전송하려면 'y' 또는 'yes'를 입력하세요 (기본값: n): ").lower().strip()
            
            if send_choice in ['y', 'yes']:
                send_email_interactive()
            else:
                print("📄 보고서 파일이 현재 폴더에 저장되었습니다.")
                
        else:
            print("❌ 보고서 생성에 실패했습니다.")
            
    except Exception as e:
        print(f"❌ 분석 중 오류가 발생했습니다: {e}")

if __name__ == "__main__":
    main()