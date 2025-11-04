"""
GitHub Actions용 판매 데이터 분석 및 보고서 생성 스크립트
원본: word_report_generator.py를 GitHub Actions 환경에 맞게 수정한 버전
"""

import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import numpy as np
import os
import warnings
warnings.filterwarnings('ignore')

# GitHub Actions 환경에서 GUI 없이 차트 생성
matplotlib.use('Agg')

# Ubuntu 환경에서 사용 가능한 폰트 설정
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

def load_and_clean_data(file_path):
    """데이터 로드 및 전처리"""
    print("=== 데이터 로드 중 ===")
    
    try:
        df = pd.read_csv(file_path)
        print(f"✅ 데이터 로드 완료: {len(df)}개")
        
        # 데이터 정리
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
        df['ProductName'] = df['ProductName'].str.title()
        df['Category'] = df['Category'].str.title()
        df['Salesperson'] = df['Salesperson'].str.title()
        
        # 결측값 및 오류 데이터 제거
        df = df.dropna(subset=['Date', 'ProductID', 'Quantity', 'UnitPrice'])
        df = df[df['ProductID'] != 'P0000']
        df = df[df['Quantity'] > 0]
        df = df[df['UnitPrice'] > 0]
        
        # TotalPrice 재계산
        df['TotalPrice'] = df['Quantity'] * df['UnitPrice']
        
        print(f"✅ 데이터 전처리 완료: {len(df)}개")
        return df
        
    except Exception as e:
        print(f"❌ 데이터 로드 실패: {e}")
        return None

def create_charts(df):
    """GitHub Actions 환경용 차트 생성"""
    print("📊 차트 생성 중...")
    chart_files = []
    
    try:
        # 1. 카테고리별 매출 파이차트
        category_sales = df.groupby('Category')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(10, 8))
        colors = plt.cm.Set3(np.linspace(0, 1, len(category_sales)))
        wedges, texts, autotexts = plt.pie(category_sales.values, labels=category_sales.index, 
                                          autopct='%1.1f%%', colors=colors, startangle=90)
        plt.title('Category Sales Distribution', fontsize=16, fontweight='bold', pad=20)
        
        # 영어로 범례 생성 (한글 폰트 이슈 방지)
        plt.legend(wedges, [f'{cat}: {val:,.0f}' for cat, val in category_sales.items()], 
                  title="Sales by Category", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        
        plt.tight_layout()
        chart1_path = 'chart_category_pie.png'
        plt.savefig(chart1_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart1_path)
        
        # 2. 지역별 매출 막대차트
        region_sales = df.groupby('Region')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(12, 8))
        bars = plt.bar(region_sales.index, region_sales.values, 
                      color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FCEA2B'])
        plt.title('Sales by Region', fontsize=16, fontweight='bold', pad=20)
        plt.xlabel('Region', fontsize=14)
        plt.ylabel('Sales Amount', fontsize=14)
        
        # 값 표시
        for bar, value in zip(bars, region_sales.values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(region_sales.values) * 0.01, 
                    f'{value:,.0f}', ha='center', va='bottom', fontsize=12, fontweight='bold')
        
        plt.xticks(rotation=45)
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        chart2_path = 'chart_region_bar.png'
        plt.savefig(chart2_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart2_path)
        
        # 3. 일별 매출 추이
        daily_sales = df.groupby('Date')['TotalPrice'].sum().sort_index()
        
        plt.figure(figsize=(14, 8))
        plt.plot(daily_sales.index, daily_sales.values, marker='o', linewidth=3, markersize=6, 
                color='#2E86AB', markerfacecolor='#F24236')
        plt.title('Daily Sales Trend', fontsize=16, fontweight='bold', pad=20)
        plt.xlabel('Date', fontsize=14)
        plt.ylabel('Sales Amount', fontsize=14)
        
        # 최고점 표시
        max_idx = daily_sales.idxmax()
        plt.annotate(f'Peak: {daily_sales[max_idx]:,.0f}', 
                    xy=(max_idx, daily_sales[max_idx]), xytext=(20, 20),
                    textcoords='offset points', ha='left',
                    bbox=dict(boxstyle='round,pad=0.5', fc='yellow', alpha=0.8),
                    arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0.2'))
        
        plt.xticks(rotation=45)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        chart3_path = 'chart_daily_trend.png'
        plt.savefig(chart3_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart3_path)
        
        print(f"✅ {len(chart_files)}개 차트 생성 완료")
        return chart_files
        
    except Exception as e:
        print(f"❌ 차트 생성 실패: {e}")
        return []

def generate_word_report(df, chart_files):
    """워드 보고서 생성 (GitHub Actions용)"""
    print("📄 워드 보고서 생성 중...")
    
    try:
        doc = Document()
        
        # 제목
        title = doc.add_heading('Sales Data Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성 정보
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        doc.add_paragraph(f"Generated by: GitHub Actions")
        doc.add_paragraph("")
        
        # 요약 통계
        doc.add_heading('📊 Summary Statistics', level=1)
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_table.cell(0, 0).text = 'Metric'
        summary_table.cell(0, 1).text = 'Value'
        
        total_sales = df['TotalPrice'].sum()
        total_quantity = df['Quantity'].sum()
        avg_order_value = df['TotalPrice'].mean()
        unique_products = df['ProductID'].nunique()
        date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
        
        summary_data = [
            ['Analysis Period', date_range],
            ['Total Sales', f"${total_sales:,.0f}"],
            ['Total Quantity', f"{total_quantity:,} units"],
            ['Average Order Value', f"${avg_order_value:,.0f}"],
            ['Product Types', f"{unique_products} products"]
        ]
        
        for i, (metric, value) in enumerate(summary_data, 1):
            summary_table.cell(i, 0).text = metric
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # 카테고리별 분석
        doc.add_heading('📊 Category Analysis', level=1)
        
        category_sales = df.groupby('Category').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum',
            'ProductID': 'nunique'
        }).round(2)
        category_sales.columns = ['Total Sales', 'Total Quantity', 'Product Count']
        category_sales = category_sales.sort_values('Total Sales', ascending=False)
        
        cat_table = doc.add_table(rows=len(category_sales)+1, cols=4)
        cat_table.style = 'Table Grid'
        
        headers = ['Category', 'Total Sales ($)', 'Total Quantity', 'Product Count']
        for i, header in enumerate(headers):
            cat_table.cell(0, i).text = header
        
        for i, (category, row) in enumerate(category_sales.iterrows(), 1):
            cat_table.cell(i, 0).text = str(category)
            cat_table.cell(i, 1).text = f"{row['Total Sales']:,.0f}"
            cat_table.cell(i, 2).text = f"{row['Total Quantity']:,.0f}"
            cat_table.cell(i, 3).text = str(row['Product Count'])
        
        # 차트 삽입
        if len(chart_files) > 0 and os.path.exists(chart_files[0]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 Category Sales Distribution Chart", style='Heading 2')
            doc.add_picture(chart_files[0], width=Inches(6))
        
        # 지역별 분석
        doc.add_paragraph("")
        doc.add_heading('🌍 Regional Analysis', level=1)
        
        region_sales = df.groupby('Region')['TotalPrice'].sum().sort_values(ascending=False)
        
        if len(chart_files) > 1 and os.path.exists(chart_files[1]):
            doc.add_paragraph("📊 Regional Sales Comparison", style='Heading 2')
            doc.add_picture(chart_files[1], width=Inches(6))
        
        # 일별 추이
        doc.add_paragraph("")
        doc.add_heading('📈 Daily Sales Trend', level=1)
        
        daily_sales_data = df.groupby('Date')['TotalPrice'].sum()
        max_sales_day = daily_sales_data.idxmax()
        min_sales_day = daily_sales_data.idxmin()
        avg_daily_sales = daily_sales_data.mean()
        
        daily_stats = doc.add_paragraph()
        daily_stats.add_run(f"• Peak Sales Day: {max_sales_day.strftime('%Y-%m-%d')} - ${daily_sales_data[max_sales_day]:,.0f}\n")
        daily_stats.add_run(f"• Lowest Sales Day: {min_sales_day.strftime('%Y-%m-%d')} - ${daily_sales_data[min_sales_day]:,.0f}\n")
        daily_stats.add_run(f"• Average Daily Sales: ${avg_daily_sales:,.0f}")
        
        if len(chart_files) > 2 and os.path.exists(chart_files[2]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 Daily Sales Trend Chart", style='Heading 2')
            doc.add_picture(chart_files[2], width=Inches(7))
        
        # 결론
        doc.add_paragraph("")
        doc.add_heading('💡 Key Insights', level=1)
        
        top_category = category_sales.index[0]
        top_region = region_sales.index[0]
        
        insights = doc.add_paragraph()
        insights.add_run(f"1. Top Category: '{top_category}' accounts for the largest portion of sales.\n\n")
        insights.add_run(f"2. Leading Region: '{top_region}' shows the highest sales performance.\n\n")
        insights.add_run(f"3. Average Daily Sales: ${avg_daily_sales:,.0f} indicates stable revenue flow.\n\n")
        insights.add_run("4. Recommendation: Focus marketing investments on top-performing categories and regions.")
        
        # 파일 저장
        report_filename = f'sales_analysis_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(report_filename)
        
        print(f"✅ 보고서 저장 완료: {report_filename}")
        return report_filename
        
    except Exception as e:
        print(f"❌ 보고서 생성 실패: {e}")
        return None

def main():
    """GitHub Actions용 메인 함수"""
    print("🚀 GitHub Actions - Sales Report Generation")
    print("="*50)
    
    try:
        # 데이터 로드
        df = load_and_clean_data('cicd_data.csv')
        if df is None:
            print("❌ 데이터 로드 실패")
            exit(1)
        
        # 차트 생성
        chart_files = create_charts(df)
        
        # 보고서 생성
        report_file = generate_word_report(df, chart_files)
        
        if report_file:
            print(f"\n✅ 보고서 생성 성공!")
            print(f"📄 파일: {report_file}")
            print(f"📊 차트: {len(chart_files)}개")
            
            # 파일 크기 정보
            if os.path.exists(report_file):
                size_mb = os.path.getsize(report_file) / (1024 * 1024)
                print(f"📦 파일 크기: {size_mb:.2f}MB")
        else:
            print("❌ 보고서 생성 실패")
            exit(1)
            
    except Exception as e:
        print(f"❌ 실행 중 오류: {e}")
        exit(1)

if __name__ == "__main__":
    main()