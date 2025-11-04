import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import numpy as np
from matplotlib import font_manager
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
import warnings
warnings.filterwarnings('ignore')

# matplotlib 백엔드 설정 (GUI 없이 이미지만 생성)
import matplotlib
matplotlib.use('Agg')

# 한글 폰트 설정 (Windows 환경에서 한글 표시를 위함)
plt.rcParams['font.family'] = 'Malgun Gothic'  # Windows 기본 한글 폰트
plt.rcParams['axes.unicode_minus'] = False  # 마이너스 기호 표시 오류 해결

def load_and_clean_data(file_path):
    """데이터 로드 및 전처리"""
    # CSV 파일 읽기
    df = pd.read_csv(file_path)
    
    print("=== 원본 데이터 정보 ===")
    print(f"총 데이터 개수: {len(df)}개")
    print(f"컬럼: {list(df.columns)}")
    print()
    
    # 데이터 정리
    # 1. 날짜 컬럼 변환
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
    
    # 2. 제품명과 카테고리 대소문자 정리
    df['ProductName'] = df['ProductName'].str.title()  # 첫 글자만 대문자
    df['Category'] = df['Category'].str.title()
    df['Salesperson'] = df['Salesperson'].str.title()
    
    # 3. 빈 값이나 잘못된 데이터 제거
    df = df.dropna(subset=['Date', 'ProductID', 'Quantity', 'UnitPrice'])
    df = df[df['ProductID'] != 'P0000']  # 잘못된 제품 ID 제거
    df = df[df['Quantity'] > 0]  # 수량이 0인 데이터 제거
    df = df[df['UnitPrice'] > 0]  # 단가가 0인 데이터 제거
    
    # 4. TotalPrice 재계산 (일부 데이터에 오류가 있을 수 있음)
    df['TotalPrice'] = df['Quantity'] * df['UnitPrice']
    
    print("=== 정리된 데이터 정보 ===")
    print(f"정리 후 데이터 개수: {len(df)}개")
    print()
    
    return df

def generate_summary_statistics(df):
    """요약 통계 생성"""
    print("="*50)
    print("📊 판매 데이터 요약 보고서")
    print("="*50)
    
    # 기본 통계
    total_sales = df['TotalPrice'].sum()
    total_quantity = df['Quantity'].sum()
    avg_order_value = df['TotalPrice'].mean()
    unique_products = df['ProductID'].nunique()
    date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
    
    print(f"📅 분석 기간: {date_range}")
    print(f"💰 총 매출액: {total_sales:,.0f}원")
    print(f"📦 총 판매 수량: {total_quantity:,}개")
    print(f"📈 평균 주문 금액: {avg_order_value:,.0f}원")
    print(f"🛍️ 판매된 제품 종류: {unique_products}개")
    print()

def analyze_by_category(df):
    """카테고리별 분석"""
    print("="*30)
    print("📊 카테고리별 판매 분석")
    print("="*30)
    
    category_sales = df.groupby('Category').agg({
        'TotalPrice': 'sum',
        'Quantity': 'sum',
        'ProductID': 'nunique'
    }).round(2)
    
    category_sales.columns = ['총 매출액', '총 수량', '제품 종류 수']
    category_sales = category_sales.sort_values('총 매출액', ascending=False)
    
    print(category_sales)
    print()
    
    return category_sales

def analyze_by_product(df):
    """제품별 분석"""
    print("="*30)
    print("🏆 베스트셀러 제품 TOP 10")
    print("="*30)
    
    product_sales = df.groupby(['ProductID', 'ProductName']).agg({
        'TotalPrice': 'sum',
        'Quantity': 'sum'
    }).round(2)
    
    product_sales.columns = ['총 매출액', '총 수량']
    top_products = product_sales.sort_values('총 매출액', ascending=False).head(10)
    
    print(top_products)
    print()
    
    return top_products

def analyze_by_region(df):
    """지역별 분석"""
    print("="*30)
    print("🌍 지역별 판매 분석")
    print("="*30)
    
    region_sales = df.groupby('Region').agg({
        'TotalPrice': 'sum',
        'Quantity': 'sum'
    }).round(2)
    
    region_sales.columns = ['총 매출액', '총 수량']
    region_sales = region_sales.sort_values('총 매출액', ascending=False)
    
    print(region_sales)
    print()
    
    return region_sales

def analyze_by_salesperson(df):
    """영업사원별 분석"""
    print("="*30)
    print("👤 영업사원별 판매 성과")
    print("="*30)
    
    # 빈 값 제거
    df_clean = df[df['Salesperson'].notna() & (df['Salesperson'] != '')]
    
    salesperson_sales = df_clean.groupby('Salesperson').agg({
        'TotalPrice': 'sum',
        'Quantity': 'sum',
        'Date': 'count'  # 거래 횟수
    }).round(2)
    
    salesperson_sales.columns = ['총 매출액', '총 수량', '거래 횟수']
    salesperson_sales = salesperson_sales.sort_values('총 매출액', ascending=False)
    
    print(salesperson_sales)
    print()
    
    return salesperson_sales

def analyze_daily_trends(df):
    """일별 판매 추이 분석"""
    print("="*30)
    print("📈 일별 판매 추이")
    print("="*30)
    
    daily_sales = df.groupby('Date').agg({
        'TotalPrice': 'sum',
        'Quantity': 'sum'
    }).round(2)
    
    daily_sales.columns = ['일별 매출액', '일별 수량']
    
    # 최고/최저 매출일
    max_sales_day = daily_sales['일별 매출액'].idxmax()
    min_sales_day = daily_sales['일별 매출액'].idxmin()
    
    print(f"최고 매출일: {max_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[max_sales_day, '일별 매출액']:,.0f}원")
    print(f"최저 매출일: {min_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[min_sales_day, '일별 매출액']:,.0f}원")
    print(f"일평균 매출액: {daily_sales['일별 매출액'].mean():,.0f}원")
    print()
    
    return daily_sales

def create_visualizations(df, category_sales, region_sales, daily_sales):
    """데이터 시각화"""
    print("="*30)
    print("📊 차트 생성 중...")
    print("="*30)
    
    # 한글 폰트 설정 재확인
    plt.rcParams['font.family'] = 'Malgun Gothic'
    plt.rcParams['axes.unicode_minus'] = False
    
    # 2x2 서브플롯 생성
    fig, axes = plt.subplots(2, 2, figsize=(15, 12))
    fig.suptitle('판매 데이터 분석 대시보드', fontsize=16, fontweight='bold')
    
    # 1. 카테고리별 매출 파이차트
    axes[0, 0].pie(category_sales['총 매출액'], labels=category_sales.index, autopct='%1.1f%%')
    axes[0, 0].set_title('카테고리별 매출 비율')
    
    # 2. 지역별 매출 막대차트
    axes[0, 1].bar(region_sales.index, region_sales['총 매출액'])
    axes[0, 1].set_title('지역별 매출액')
    axes[0, 1].set_ylabel('매출액 (원)')
    for i, v in enumerate(region_sales['총 매출액']):
        axes[0, 1].text(i, v + max(region_sales['총 매출액']) * 0.01, f'{v:,.0f}', ha='center')
    
    # 3. 일별 매출 추이 선 그래프
    axes[1, 0].plot(daily_sales.index, daily_sales['일별 매출액'], marker='o')
    axes[1, 0].set_title('일별 매출 추이')
    axes[1, 0].set_ylabel('매출액 (원)')
    axes[1, 0].tick_params(axis='x', rotation=45)
    
    # 4. 제품별 수량 TOP 5 막대차트
    top5_products = df.groupby('ProductName')['Quantity'].sum().sort_values(ascending=False).head(5)
    axes[1, 1].barh(range(len(top5_products)), top5_products.values)
    axes[1, 1].set_yticks(range(len(top5_products)))
    axes[1, 1].set_yticklabels(top5_products.index)
    axes[1, 1].set_title('TOP 5 판매량 제품')
    axes[1, 1].set_xlabel('판매 수량')
    
    plt.tight_layout()
    plt.savefig('sales_analysis_dashboard.png', dpi=300, bbox_inches='tight')
    print("차트가 'sales_analysis_dashboard.png' 파일로 저장되었습니다.")
    plt.show()

def generate_excel_report(df, category_sales, region_sales, salesperson_sales, daily_sales):
    """Excel 보고서 생성"""
    print("="*30)
    print("📝 Excel 보고서 생성 중...")
    print("="*30)
    
    with pd.ExcelWriter('sales_analysis_report.xlsx', engine='openpyxl') as writer:
        # 원본 데이터 (정리된 버전)
        df.to_excel(writer, sheet_name='원본데이터', index=False)
        
        # 요약 통계
        summary_data = {
            '구분': ['총 매출액', '총 판매수량', '평균 주문금액', '제품 종류 수', '분석 기간'],
            '값': [
                f"{df['TotalPrice'].sum():,.0f}원",
                f"{df['Quantity'].sum():,}개",
                f"{df['TotalPrice'].mean():,.0f}원",
                f"{df['ProductID'].nunique()}개",
                f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
            ]
        }
        pd.DataFrame(summary_data).to_excel(writer, sheet_name='요약통계', index=False)
        
        # 각 분석 결과를 별도 시트에 저장
        category_sales.to_excel(writer, sheet_name='카테고리별분석')
        region_sales.to_excel(writer, sheet_name='지역별분석')
        salesperson_sales.to_excel(writer, sheet_name='영업사원별분석')
        daily_sales.to_excel(writer, sheet_name='일별추이')
        
        # 베스트셀러 제품
        top_products = df.groupby(['ProductID', 'ProductName']).agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).sort_values('TotalPrice', ascending=False).head(10)
        top_products.to_excel(writer, sheet_name='베스트셀러제품')
    
    print("Excel 보고서가 'sales_analysis_report.xlsx' 파일로 저장되었습니다.")

def generate_word_report(df, category_sales, region_sales, salesperson_sales, daily_sales):
    """워드 파일(.docx) 보고서 생성"""
    print("="*30)
    print("📄 Word 보고서 생성 중...")
    print("="*30)
    
    try:
        # 새 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading('판매 데이터 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성 일자 추가
        doc.add_paragraph(f"보고서 생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        doc.add_paragraph("")  # 빈 줄
        
        # === 1. 요약 통계 ===
        doc.add_heading('📊 요약 통계', level=1)
        
        # 요약 통계 표 생성
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        # 표 헤더
        summary_table.cell(0, 0).text = '구분'
        summary_table.cell(0, 1).text = '값'
        
        # 요약 데이터
        total_sales = df['TotalPrice'].sum()
        total_quantity = df['Quantity'].sum()
        avg_order_value = df['TotalPrice'].mean()
        unique_products = df['ProductID'].nunique()
        date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
        
        summary_data = [
            ['분석 기간', date_range],
            ['총 매출액', f"{total_sales:,.0f}원"],
            ['총 판매 수량', f"{total_quantity:,}개"],
            ['평균 주문 금액', f"{avg_order_value:,.0f}원"],
            ['판매된 제품 종류', f"{unique_products}개"]
        ]
        
        for i, (category, value) in enumerate(summary_data, 1):
            summary_table.cell(i, 0).text = category
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 2. 카테고리별 분석 ===
        doc.add_heading('📊 카테고리별 판매 분석', level=1)
        
        # 카테고리별 분석 표
        cat_table = doc.add_table(rows=len(category_sales)+1, cols=4)
        cat_table.style = 'Table Grid'
        
        # 표 헤더
        cat_headers = ['카테고리', '총 매출액 (원)', '총 수량 (개)', '제품 종류 수']
        for i, header in enumerate(cat_headers):
            cat_table.cell(0, i).text = header
        
        # 카테고리 데이터
        for i, (category, row) in enumerate(category_sales.iterrows(), 1):
            cat_table.cell(i, 0).text = category
            cat_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            cat_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
            cat_table.cell(i, 3).text = str(row['제품 종류 수'])
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 3. 지역별 분석 ===
        doc.add_heading('🌍 지역별 판매 분석', level=1)
        
        # 지역별 분석 표
        region_table = doc.add_table(rows=len(region_sales)+1, cols=3)
        region_table.style = 'Table Grid'
        
        # 표 헤더
        region_headers = ['지역', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(region_headers):
            region_table.cell(0, i).text = header
        
        # 지역 데이터
        for i, (region, row) in enumerate(region_sales.iterrows(), 1):
            region_table.cell(i, 0).text = region
            region_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            region_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 4. 베스트셀러 제품 ===
        doc.add_heading('🏆 베스트셀러 제품 TOP 10', level=1)
        
        # 베스트셀러 제품 데이터
        top_products = df.groupby(['ProductID', 'ProductName']).agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).sort_values('TotalPrice', ascending=False).head(10)
        
        # 베스트셀러 표
        product_table = doc.add_table(rows=len(top_products)+1, cols=4)
        product_table.style = 'Table Grid'
        
        # 표 헤더
        product_headers = ['순위', '제품명', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(product_headers):
            product_table.cell(0, i).text = header
        
        # 제품 데이터
        for i, ((product_id, product_name), row) in enumerate(top_products.iterrows(), 1):
            product_table.cell(i, 0).text = str(i)
            product_table.cell(i, 1).text = product_name
            product_table.cell(i, 2).text = f"{row['TotalPrice']:,.0f}"
            product_table.cell(i, 3).text = f"{row['Quantity']:,.0f}"
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 5. 영업사원별 분석 ===
        doc.add_heading('👤 영업사원별 판매 성과', level=1)
        
        # 빈 값 제거
        df_clean = df[df['Salesperson'].notna() & (df['Salesperson'] != '')]
        salesperson_data = df_clean.groupby('Salesperson').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum',
            'Date': 'count'
        }).sort_values('TotalPrice', ascending=False)
        
        # 영업사원 분석 표
        sales_table = doc.add_table(rows=len(salesperson_data)+1, cols=4)
        sales_table.style = 'Table Grid'
        
        # 표 헤더
        sales_headers = ['영업사원', '총 매출액 (원)', '총 수량 (개)', '거래 횟수']
        for i, header in enumerate(sales_headers):
            sales_table.cell(0, i).text = header
        
        # 영업사원 데이터
        for i, (salesperson, row) in enumerate(salesperson_data.iterrows(), 1):
            sales_table.cell(i, 0).text = salesperson
            sales_table.cell(i, 1).text = f"{row['TotalPrice']:,.0f}"
            sales_table.cell(i, 2).text = f"{row['Quantity']:,.0f}"
            sales_table.cell(i, 3).text = str(row['Date'])
        
        doc.add_paragraph("")  # 빈 줄
        
        # === 6. 일별 매출 추이 분석 ===
        doc.add_heading('📈 일별 매출 추이 분석', level=1)
        
        # 일별 통계
        max_sales_day = daily_sales['일별 매출액'].idxmax()
        min_sales_day = daily_sales['일별 매출액'].idxmin()
        avg_daily_sales = daily_sales['일별 매출액'].mean()
        
        daily_stats = doc.add_paragraph()
        daily_stats.add_run(f"• 최고 매출일: {max_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[max_sales_day, '일별 매출액']:,.0f}원\n")
        daily_stats.add_run(f"• 최저 매출일: {min_sales_day.strftime('%Y-%m-%d')} - {daily_sales.loc[min_sales_day, '일별 매출액']:,.0f}원\n")
        daily_stats.add_run(f"• 일평균 매출액: {avg_daily_sales:,.0f}원")
        
        # 일별 매출 TOP 10
        doc.add_paragraph("")
        doc.add_heading('일별 매출 TOP 10', level=2)
        
        top10_days = daily_sales.sort_values('일별 매출액', ascending=False).head(10)
        
        daily_table = doc.add_table(rows=len(top10_days)+1, cols=3)
        daily_table.style = 'Table Grid'
        
        # 표 헤더
        daily_headers = ['순위', '날짜', '매출액 (원)']
        for i, header in enumerate(daily_headers):
            daily_table.cell(0, i).text = header
        
        # 일별 데이터
        for i, (date, row) in enumerate(top10_days.iterrows(), 1):
            daily_table.cell(i, 0).text = str(i)
            daily_table.cell(i, 1).text = date.strftime('%Y-%m-%d')
            daily_table.cell(i, 2).text = f"{row['일별 매출액']:,.0f}"
        
        # 문서 저장
        doc.save('sales_analysis_report.docx')
        print("✅ Word 보고서가 'sales_analysis_report.docx' 파일로 저장되었습니다.")
        
    except Exception as e:
        print(f"❌ Word 파일 생성 중 오류 발생: {e}")
        print("python-docx 라이브러리가 제대로 설치되었는지 확인해주세요.")

def main():
    """메인 함수"""
    print("🚀 판매 데이터 분석을 시작합니다...\n")
    
    # 데이터 로드 및 전처리
    df = load_and_clean_data('cicd_data.csv')
    
    # 각종 분석 수행
    generate_summary_statistics(df)
    category_sales = analyze_by_category(df)
    analyze_by_product(df)
    region_sales = analyze_by_region(df)
    salesperson_sales = analyze_by_salesperson(df)
    daily_sales = analyze_daily_trends(df)
    
    # 시각화 생성
    create_visualizations(df, category_sales, region_sales, daily_sales)
    
    # Excel 보고서 생성
    generate_excel_report(df, category_sales, region_sales, salesperson_sales, daily_sales)
    
    # Word 보고서 생성
    generate_word_report(df, category_sales, region_sales, salesperson_sales, daily_sales)
    
    print("\n" + "="*50)
    print("✅ 분석이 완료되었습니다!")
    print("📊 생성된 파일:")
    print("   - sales_analysis_dashboard.png (시각화 차트)")
    print("   - sales_analysis_report.xlsx (Excel 보고서)")
    print("   - sales_analysis_report.docx (Word 보고서)")
    print("="*50)

if __name__ == "__main__":
    main()