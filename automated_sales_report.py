"""
판매 데이터 분석 및 자동 보고서 시스템
CSV 파일 읽기 → 데이터 분석 → 워드 보고서 생성 → 이메일 전송
"""

import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import numpy as np
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import warnings
warnings.filterwarnings('ignore')

# matplotlib 백엔드 설정 (GUI 없이 이미지만 생성)
matplotlib.use('Agg')

# 한글 폰트 설정 (Windows 환경)
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False

# ==================== 이메일 설정 (임시) ====================
EMAIL_CONFIG = {
    'sender_email': 'your_email@gmail.com',  # 보내는 사람 이메일
    'sender_password': 'your_app_password',  # Gmail 앱 비밀번호
    'recipient_emails': [                    # 받는 사람 이메일 리스트
        'recipient1@gmail.com',
        'recipient2@company.com'
    ],
    'smtp_server': 'smtp.gmail.com',         # SMTP 서버
    'smtp_port': 587                         # SMTP 포트
}

def load_and_clean_data(file_path):
    """CSV 데이터 로드 및 전처리"""
    print("="*50)
    print("📊 데이터 로드 및 전처리")
    print("="*50)
    
    try:
        # CSV 파일 읽기
        df = pd.read_csv(file_path, encoding='utf-8')
        print(f"✅ 데이터 로드 완료: {file_path}")
        print(f"📊 원본 데이터 개수: {len(df)}개")
        print(f"📋 컬럼: {list(df.columns)}")
        
        # 데이터 정리
        print("\n🔧 데이터 전처리 중...")
        
        # 1. 날짜 컬럼 변환
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
        
        # 2. 텍스트 데이터 정리
        df['ProductName'] = df['ProductName'].str.title()
        df['Category'] = df['Category'].str.title()
        df['Salesperson'] = df['Salesperson'].str.title()
        
        # 3. 결측값 및 오류 데이터 제거
        original_count = len(df)
        df = df.dropna(subset=['Date', 'ProductID', 'Quantity', 'UnitPrice'])
        df = df[df['ProductID'] != 'P0000']  # 잘못된 제품 ID 제거
        df = df[df['Quantity'] > 0]
        df = df[df['UnitPrice'] > 0]
        
        # 4. TotalPrice 재계산
        df['TotalPrice'] = df['Quantity'] * df['UnitPrice']
        
        cleaned_count = len(df)
        removed_count = original_count - cleaned_count
        
        print(f"✅ 데이터 전처리 완료")
        print(f"📊 정리 후 데이터: {cleaned_count}개")
        print(f"🗑️  제거된 데이터: {removed_count}개")
        
        return df
        
    except FileNotFoundError:
        print(f"❌ 파일을 찾을 수 없습니다: {file_path}")
        return None
    except Exception as e:
        print(f"❌ 데이터 로드 중 오류 발생: {e}")
        return None

def create_charts(df):
    """시각화 차트 생성"""
    print("\n📊 시각화 차트 생성 중...")
    
    chart_files = []
    
    try:
        # 1. 카테고리별 매출 파이차트
        category_sales = df.groupby('Category')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(8, 6))
        colors = plt.cm.Set3(np.linspace(0, 1, len(category_sales)))
        wedges, texts, autotexts = plt.pie(category_sales.values, labels=category_sales.index, 
                                          autopct='%1.1f%%', colors=colors, startangle=90)
        plt.title('카테고리별 매출 비율', fontsize=14, fontweight='bold', pad=20)
        
        plt.legend(wedges, [f'{cat}\n{val:,.0f}원' for cat, val in category_sales.items()], 
                  title="카테고리별 매출액", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        
        plt.tight_layout()
        chart1_path = 'chart_category_pie.png'
        plt.savefig(chart1_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart1_path)
        
        # 2. 지역별 매출 막대차트
        region_sales = df.groupby('Region')['TotalPrice'].sum().sort_values(ascending=False)
        
        plt.figure(figsize=(10, 6))
        bars = plt.bar(region_sales.index, region_sales.values, color=['#FF9999', '#66B2FF', '#99FF99', '#FFCC99'])
        plt.title('지역별 매출액', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('지역', fontsize=12)
        plt.ylabel('매출액 (원)', fontsize=12)
        
        for bar, value in zip(bars, region_sales.values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(region_sales.values) * 0.01, 
                    f'{value:,.0f}원', ha='center', va='bottom', fontsize=10)
        
        plt.xticks(rotation=45)
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        chart2_path = 'chart_region_bar.png'
        plt.savefig(chart2_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart2_path)
        
        # 3. 일별 매출 추이 선 그래프
        daily_sales = df.groupby('Date')['TotalPrice'].sum().sort_index()
        
        plt.figure(figsize=(12, 6))
        plt.plot(daily_sales.index, daily_sales.values, marker='o', linewidth=2, markersize=4, color='#2E86AB')
        plt.title('일별 매출 추이', fontsize=14, fontweight='bold', pad=20)
        plt.xlabel('날짜', fontsize=12)
        plt.ylabel('매출액 (원)', fontsize=12)
        
        max_idx = daily_sales.idxmax()
        plt.annotate(f'최고: {daily_sales[max_idx]:,.0f}원', 
                    xy=(max_idx, daily_sales[max_idx]), xytext=(10, 10),
                    textcoords='offset points', ha='left',
                    bbox=dict(boxstyle='round,pad=0.3', fc='yellow', alpha=0.7),
                    arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0'))
        
        plt.xticks(rotation=45)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        chart3_path = 'chart_daily_trend.png'
        plt.savefig(chart3_path, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart3_path)
        
        print(f"✅ {len(chart_files)}개의 차트가 생성되었습니다.")
        return chart_files
        
    except Exception as e:
        print(f"❌ 차트 생성 중 오류 발생: {e}")
        return []

def generate_word_report(df, chart_files):
    """워드 보고서 생성"""
    print("\n📄 워드 보고서 생성 중...")
    
    try:
        # 새 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading('판매 데이터 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성 일자 및 기본 정보
        doc.add_paragraph(f"보고서 생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}")
        doc.add_paragraph(f"데이터 분석 기간: {df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}")
        doc.add_paragraph("")
        
        # === 1. 요약 통계 ===
        doc.add_heading('📊 요약 통계', level=1)
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_table.cell(0, 0).text = '구분'
        summary_table.cell(0, 1).text = '값'
        
        total_sales = df['TotalPrice'].sum()
        total_quantity = df['Quantity'].sum()
        avg_order_value = df['TotalPrice'].mean()
        unique_products = df['ProductID'].nunique()
        date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} ~ {df['Date'].max().strftime('%Y-%m-%d')}"
        
        summary_data = [
            ['분석 기간', date_range],
            ['총 매출액', f"{total_sales:,.0f}원"],
            ['총 판매 수량', f"{total_quantity:,}개"],
            ['평균 주문 금액', f"{avg_order_value:,.0f}원"],
            ['판매된 제품 종류', f"{unique_products}개"]
        ]
        
        for i, (category, value) in enumerate(summary_data, 1):
            summary_table.cell(i, 0).text = category
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # === 2. 카테고리별 분석 ===
        doc.add_heading('📊 카테고리별 판매 분석', level=1)
        
        category_sales = df.groupby('Category').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum',
            'ProductID': 'nunique'
        }).round(2)
        category_sales.columns = ['총 매출액', '총 수량', '제품 종류 수']
        category_sales = category_sales.sort_values('총 매출액', ascending=False)
        
        cat_table = doc.add_table(rows=len(category_sales)+1, cols=4)
        cat_table.style = 'Table Grid'
        
        cat_headers = ['카테고리', '총 매출액 (원)', '총 수량 (개)', '제품 종류 수']
        for i, header in enumerate(cat_headers):
            cat_table.cell(0, i).text = header
        
        for i, (category, row) in enumerate(category_sales.iterrows(), 1):
            cat_table.cell(i, 0).text = category
            cat_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            cat_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
            cat_table.cell(i, 3).text = str(row['제품 종류 수'])
        
        # 카테고리 차트 삽입
        if len(chart_files) > 0 and os.path.exists(chart_files[0]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 카테고리별 매출 비율 차트", style='Heading 2')
            doc.add_picture(chart_files[0], width=Inches(6))
        
        doc.add_paragraph("")
        
        # === 3. 지역별 분석 ===
        doc.add_heading('🌍 지역별 판매 분석', level=1)
        
        region_sales = df.groupby('Region').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).round(2)
        region_sales.columns = ['총 매출액', '총 수량']
        region_sales = region_sales.sort_values('총 매출액', ascending=False)
        
        region_table = doc.add_table(rows=len(region_sales)+1, cols=3)
        region_table.style = 'Table Grid'
        
        region_headers = ['지역', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(region_headers):
            region_table.cell(0, i).text = header
        
        for i, (region, row) in enumerate(region_sales.iterrows(), 1):
            region_table.cell(i, 0).text = region
            region_table.cell(i, 1).text = f"{row['총 매출액']:,.0f}"
            region_table.cell(i, 2).text = f"{row['총 수량']:,.0f}"
        
        # 지역별 차트 삽입
        if len(chart_files) > 1 and os.path.exists(chart_files[1]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 지역별 매출 비교 차트", style='Heading 2')
            doc.add_picture(chart_files[1], width=Inches(6))
        
        doc.add_paragraph("")
        
        # === 4. 베스트셀러 제품 ===
        doc.add_heading('🏆 베스트셀러 제품 TOP 10', level=1)
        
        top_products = df.groupby(['ProductID', 'ProductName']).agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).sort_values('TotalPrice', ascending=False).head(10)
        
        product_table = doc.add_table(rows=len(top_products)+1, cols=4)
        product_table.style = 'Table Grid'
        
        product_headers = ['순위', '제품명', '총 매출액 (원)', '총 수량 (개)']
        for i, header in enumerate(product_headers):
            product_table.cell(0, i).text = header
        
        for i, ((product_id, product_name), row) in enumerate(top_products.iterrows(), 1):
            product_table.cell(i, 0).text = str(i)
            product_table.cell(i, 1).text = product_name
            product_table.cell(i, 2).text = f"{row['TotalPrice']:,.0f}"
            product_table.cell(i, 3).text = f"{row['Quantity']:,.0f}"
        
        doc.add_paragraph("")
        
        # === 5. 일별 매출 추이 ===
        doc.add_heading('📈 일별 매출 추이', level=1)
        
        daily_sales_data = df.groupby('Date')['TotalPrice'].sum()
        max_sales_day = daily_sales_data.idxmax()
        min_sales_day = daily_sales_data.idxmin()
        avg_daily_sales = daily_sales_data.mean()
        
        daily_stats = doc.add_paragraph()
        daily_stats.add_run(f"• 최고 매출일: {max_sales_day.strftime('%Y-%m-%d')} - {daily_sales_data[max_sales_day]:,.0f}원\n")
        daily_stats.add_run(f"• 최저 매출일: {min_sales_day.strftime('%Y-%m-%d')} - {daily_sales_data[min_sales_day]:,.0f}원\n")
        daily_stats.add_run(f"• 일평균 매출액: {avg_daily_sales:,.0f}원")
        
        # 일별 추이 차트 삽입
        if len(chart_files) > 2 and os.path.exists(chart_files[2]):
            doc.add_paragraph("")
            doc.add_paragraph("📊 일별 매출 추이 차트", style='Heading 2')
            doc.add_picture(chart_files[2], width=Inches(7))
        
        # === 6. 결론 및 제안사항 ===
        doc.add_heading('💡 결론 및 제안사항', level=1)
        
        top_category = category_sales.index[0]
        top_region = region_sales.index[0]
        
        conclusions = doc.add_paragraph()
        conclusions.add_run(f"1. 핵심 카테고리: '{top_category}' 카테고리가 전체 매출의 주요 부분을 차지합니다.\n\n")
        conclusions.add_run(f"2. 주요 지역: '{top_region}' 지역에서 가장 높은 매출을 기록했습니다.\n\n")
        conclusions.add_run(f"3. 일평균 매출: {avg_daily_sales:,.0f}원으로 안정적인 매출을 보이고 있습니다.\n\n")
        conclusions.add_run("4. 제안사항: 상위 성과 카테고리와 지역에 대한 집중적인 마케팅 투자를 고려해보시기 바랍니다.")
        
        # 문서 저장
        report_filename = f'sales_analysis_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(report_filename)
        
        print(f"✅ 워드 보고서 생성 완료: {report_filename}")
        return report_filename
        
    except Exception as e:
        print(f"❌ 워드 보고서 생성 중 오류 발생: {e}")
        return None

def send_email_with_report(report_file, config):
    """이메일로 보고서 전송"""
    print(f"\n📧 이메일 전송 중...")
    
    try:
        # 이메일 메시지 생성
        msg = MIMEMultipart()
        msg['From'] = config['sender_email']
        msg['To'] = ", ".join(config['recipient_emails'])
        msg['Subject'] = f"판매 데이터 분석 보고서 - {datetime.now().strftime('%Y년 %m월 %d일')}"
        
        # 이메일 본문
        body = f"""
안녕하세요,

첨부된 파일은 {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}에 자동 생성된 판매 데이터 분석 보고서입니다.

📊 보고서 주요 내용:
• 매출 요약 통계
• 카테고리별 판매 분석 (차트 포함)
• 지역별 성과 비교 (차트 포함)
• 베스트셀러 제품 순위
• 일별 매출 추이 분석 (차트 포함)
• 결론 및 제안사항

📈 데이터 기반 인사이트와 시각화 차트가 포함되어 있으니 검토 후 피드백 부탁드립니다.

감사합니다.

---
⚡ 자동화 시스템으로 생성된 보고서
🕒 생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        # 본문 첨부
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        # 워드 파일 첨부
        if os.path.exists(report_file):
            with open(report_file, "rb") as attachment:
                part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            filename = os.path.basename(report_file)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {filename}',
            )
            msg.attach(part)
            
            print(f"📎 파일 첨부 완료: {filename}")
        else:
            print(f"❌ 보고서 파일을 찾을 수 없습니다: {report_file}")
            return False
        
        # SMTP 서버 연결 및 전송
        server = smtplib.SMTP(config['smtp_server'], config['smtp_port'])
        server.starttls()
        server.login(config['sender_email'], config['sender_password'])
        
        text = msg.as_string()
        server.sendmail(config['sender_email'], config['recipient_emails'], text)
        server.quit()
        
        print("✅ 이메일 전송 완료!")
        print(f"📧 받는 사람: {', '.join(config['recipient_emails'])}")
        
        return True
        
    except smtplib.SMTPAuthenticationError:
        print("❌ 이메일 인증 실패 - 이메일 주소와 앱 비밀번호를 확인하세요.")
        return False
    except Exception as e:
        print(f"❌ 이메일 전송 중 오류 발생: {e}")
        return False

def main():
    """메인 함수 - 전체 프로세스 실행"""
    print("🚀 판매 데이터 분석 및 자동 보고서 시스템")
    print("="*60)
    print("📋 프로세스: CSV 읽기 → 분석 → 워드 보고서 → 이메일 전송")
    print("="*60)
    
    start_time = datetime.now()
    
    try:
        # Step 1: CSV 데이터 로드 및 전처리
        csv_file = 'cicd_data.csv'
        df = load_and_clean_data(csv_file)
        
        if df is None:
            print("❌ 데이터 로드에 실패했습니다. 프로그램을 종료합니다.")
            return
        
        # Step 2: 시각화 차트 생성
        chart_files = create_charts(df)
        
        # Step 3: 워드 보고서 생성
        report_file = generate_word_report(df, chart_files)
        
        if report_file is None:
            print("❌ 보고서 생성에 실패했습니다.")
            return
        
        # Step 4: 이메일 전송
        print(f"\n📧 이메일 설정 확인...")
        print(f"보내는 사람: {EMAIL_CONFIG['sender_email']}")
        print(f"받는 사람: {', '.join(EMAIL_CONFIG['recipient_emails'])}")
        
        # 실제 이메일 전송 (주석 해제하여 사용)
        # email_success = send_email_with_report(report_file, EMAIL_CONFIG)
        
        # 테스트용 - 실제로는 위 라인을 주석 해제하고 아래 라인을 주석 처리
        print("📧 이메일 전송 기능은 설정 완료 후 사용하세요.")
        print("💡 EMAIL_CONFIG에서 실제 이메일 정보를 입력하고 send_email_with_report 함수 주석을 해제하세요.")
        email_success = True  # 테스트용
        
        # 결과 요약
        end_time = datetime.now()
        execution_time = (end_time - start_time).seconds
        
        print("\n" + "="*60)
        print("🎉 자동 보고서 시스템 실행 완료!")
        print("="*60)
        print(f"📊 분석된 데이터: {len(df)}건")
        print(f"📄 생성된 보고서: {report_file}")
        print(f"📈 생성된 차트: {len(chart_files)}개")
        print(f"📧 이메일 전송: {'성공' if email_success else '실패'}")
        print(f"⏱️  총 실행 시간: {execution_time}초")
        print("="*60)
        
        # 생성된 파일 목록 출력
        print("\n📁 생성된 파일 목록:")
        if os.path.exists(report_file):
            file_size = os.path.getsize(report_file) / 1024  # KB
            print(f"   📄 {report_file} ({file_size:.1f}KB)")
        
        for chart_file in chart_files:
            if os.path.exists(chart_file):
                file_size = os.path.getsize(chart_file) / 1024  # KB
                print(f"   📊 {chart_file} ({file_size:.1f}KB)")
        
        print(f"\n💡 보고서를 확인하려면 '{report_file}' 파일을 열어보세요!")
        
    except KeyboardInterrupt:
        print("\n❌ 사용자가 프로그램을 중단했습니다.")
    except Exception as e:
        print(f"\n❌ 예상치 못한 오류가 발생했습니다: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()